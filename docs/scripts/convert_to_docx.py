import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]

def process_inline(paragraph, text):
    # Process links [text](url) -> just text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    parts = text.split('**')
    for i, part in enumerate(parts):
        # Handling `backticks` for code styling
        sub_parts = part.split('`')
        for j, sub_part in enumerate(sub_parts):
            if sub_part == "": continue
            run = paragraph.add_run(sub_part)
            if i % 2 == 1:
                run.bold = True
            if j % 2 == 1:
                run.font.name = 'Courier New'

def md_to_docx(md_path, docx_path):
    md_path = Path(md_path)
    docx_path = Path(docx_path)
    if not md_path.exists():
        print(f"File not found: {md_path}")
        return
        
    # We create a new Document for each file
    doc = Document()
    
    lines = md_path.read_text(encoding='utf-8').splitlines()
        
    in_code_block = False
    
    for line in lines:
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            p.style = 'No Spacing'
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            continue
            
        header_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            level = min(level, 9)
            h = doc.add_heading(level=level if level > 0 else 1)
            process_inline(h, text)
            continue
            
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', line)
        if list_match:
            text = list_match.group(3)
            if list_match.group(2) in ['-', '*']:
                p = doc.add_paragraph(style='List Bullet')
                process_inline(p, text)
            else:
                p = doc.add_paragraph(style='List Number')
                process_inline(p, text)
            continue
            
        if line.strip() == '':
            continue
            
        p = doc.add_paragraph()
        process_inline(p, line)
        
    doc.save(docx_path)
    print(f"Created {docx_path}")

try:
    md_to_docx(ROOT / 'ARCH.md', ROOT / 'Project_Architecture.docx')
    md_to_docx(ROOT / 'INSTALL.md', ROOT / 'Installation_Configuration.docx')
except Exception as e:
    print(f"Error: {e}")
