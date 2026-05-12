import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]

# Same helper as before
def process_inline(paragraph, text):
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    parts = text.split('**')
    for i, part in enumerate(parts):
        sub_parts = part.split('`')
        for j, sub_part in enumerate(sub_parts):
            if sub_part == "": continue
            run = paragraph.add_run(sub_part)
            if i % 2 == 1:
                run.bold = True
            if j % 2 == 1:
                run.font.name = 'Courier New'


content = """# Architecture du Projet Procurement Watch

## Aperçu Général
Cette application est une plateforme de renseignement sur les appels d'offres (procurement) comprenant trois parties principales :
1. Un **frontend React** utilisé pour l'authentification, la revue des projets, les commentaires, les outils administrateur et les contrôles de synchronisation.
2. Un **backend FastAPI** contenant les API, l'authentification, les uploads de fichiers, les planifications (schedulers) et les envois de notifications en direct.
3. Un **pipeline IA et de collecte de données (Scraping)** qui glane des offres sur de nombreux sites, vérifie leur domaine via une intelligence artificielle, et les insère dans MongoDB.

## Schéma Haut Niveau

```text
Frontend (React)
  -> API FastAPI
    -> MongoDB
    -> Téléchargements de fichiers (Uploads)
    -> APScheduler (pour la planification)
    -> Sous-processus de Synchronisation (main.py)
         -> Extracteurs (Scrapers de sites tiers)
         -> Vérification par Intelligence Artificielle (Filtre)
         -> Enrichissement IA (Détails approfondis)
         -> Normalisation / Déduplication / Sauvegarde des résultats
```

## Structure du Dépôt

### À la racine :
- `frontend/` : Application React / Vite.
- `backend/` : Application FastAPI, accès données, orchestrateur, scrapers.
- `docker-compose.yml` : Configuration réseau des conteneurs pour exécuter l'ensemble.
- `INSTALL.md` : Guide d'installation et déploiement.
- `ARCH.md` : Document dont ce dossier s'inspire.

## Architecture Frontend (React)

Points critiques de l'application cliente :
- `frontend/src/main.jsx` et `App.jsx` sont les piliers.
- Ils gèrent ensemble le routage, la session sécurisée et les variables globales de composant.

### Principales Interfaces :
- Tableau de bord des résultats d'offres (Avec tri par colonne, filtres rapides).
- Inspecteur Latéral (Le panneau s'ouvrant avec le détail d'un projet).
- Modules de synchronisation et de planification automatique.
- Interface d'administration pour la gestion des rôles.

## Architecture Backend (FastAPI Python)

### Les fichiers à retenir :
- `backend/server.py` : Entrée principale réseau. Fournit des réponses HTTP, des websockets (SSE) et connecte la planification APScheduler.
- `backend/main.py` : Pilote les tâches longues. Dès qu'une synchronisation démarre, c'est ce nœud qui déclenche les `scrapers` et les appels à l'API `OpenAI/DeepSeek`.
- `backend/database.py` : Gère de façon unifiée la persistance sur MongoDB.
- `backend/ai_filter.py` et `backend/ai_enrichment.py` : Moteurs de compréhension IA des descriptions de projets pour décider si ceux-ci correspondent au domaine cybersécurité (ou autres).

## Architectures de Synchronisation (Scraping)

Le cœur de ce projet repose sur le "Sourcing" automatisé.
Voici les étapes par lesquelles passent les données :

1. Les sources sélectionnées (comme DGMarket, World Bank, etc.) sont scrutées par les `scrapers` en parallèle.
2. Les données récoltées sont fusionnées et dédupliquées selon leurs identifiants internes et textes similaires.
3. Comparaison MongoDB : Le système repère ce qui existe déjà pour ne récupérer que les nouveaux projets bruts.
4. Vérification IA (Filtre) : Chaque projet passe un test "Pass/Fail" via l'IA pour économiser sur l'exploitation des projets inutiles.
5. Enrichissement IA : Si l'offre passe l'étape 4, elle est envoyée à un algorithme plus poussé pour générer un contexte résumé approfondi.
6. Sauvegarde en Base : Tout est persifié, et un fichier de rapport `.xlsx` est régénéré automatiquement.

## Commentaires et Pièces Jointes

Chaque projet incorpore un fil de discussion.
Les fichiers sont stockés dans le répertoire physique : `backend/uploads/`.
Le frontend met en œuvre un lecteur PDF en direct intégratif et compresse visuellement les galeries d'images pour faciliter la détection de besoins, par les analystes réseau.
"""

doc_path = ROOT / "Architecture_Projet_FR.docx"
doc = Document()
doc.styles['Normal'].font.name = 'Arial'

in_code_block = False

for line in content.split('\n'):
    if line.startswith('```'):
        in_code_block = not in_code_block
        continue
        
    if in_code_block:
        p = doc.add_paragraph()
        p.style = 'No Spacing'
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        continue
        
    header_match = re.match(r'^(#{1,6})\s+(.*)', line)
    if header_match:
        level = len(header_match.group(1))
        text = header_match.group(2)
        h = doc.add_heading(level=min(level, 9) if level > 0 else 1)
        process_inline(h, text)
        continue
        
    list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', line)
    if list_match:
        text = list_match.group(3)
        if list_match.group(2) in ['-', '*']:
            p = doc.add_paragraph(style='List Bullet')
            process_inline(p, text)
        else:
            p = doc.add_paragraph(style='List Number')
            process_inline(p, text)
        continue
        
    if line.strip() == '':
        continue
        
    p = doc.add_paragraph()
    process_inline(p, line)

doc.save(doc_path)
print(f"Created {doc_path}")
