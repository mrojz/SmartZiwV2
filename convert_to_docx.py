import os
import re
from docx import Document
from docx.shared import Pt

def process_inline(paragraph, text):
    # Process links [text](url) -> just text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    parts = text.split('**')
    for i, part in enumerate(parts):
        # Handling `backticks` for code styling
        sub_parts = part.split('`')
        for j, sub_part in enumerate(sub_parts):
            if sub_part == "": continue
            run = paragraph.add_run(sub_part)
            if i % 2 == 1:
                run.bold = True
            if j % 2 == 1:
                run.font.name = 'Courier New'

def md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"File not found: {md_path}")
        return
        
    # We create a new Document for each file
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    
    for line in lines:
        line = line.replace('\n', '')
        
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            p.style = 'No Spacing'
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            continue
            
        header_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            level = min(level, 9)
            h = doc.add_heading(level=level if level > 0 else 1)
            process_inline(h, text)
            continue
            
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', line)
        if list_match:
            text = list_match.group(3)
            if list_match.group(2) in ['-', '*']:
                p = doc.add_paragraph(style='List Bullet')
                process_inline(p, text)
            else:
                p = doc.add_paragraph(style='List Number')
                process_inline(p, text)
            continue
            
        if line.strip() == '':
            continue
            
        p = doc.add_paragraph()
        process_inline(p, line)
        
    doc.save(docx_path)
    print(f"Created {docx_path}")

try:
    md_to_docx(r'd:\Dev\Ziw\new_cdx_gpt_5.4\ARCH.md', r'd:\Dev\Ziw\new_cdx_gpt_5.4\Project_Architecture.docx')
    md_to_docx(r'd:\Dev\Ziw\new_cdx_gpt_5.4\INSTALL.md', r'd:\Dev\Ziw\new_cdx_gpt_5.4\Installation_Configuration.docx')
except Exception as e:
    print(f"Error: {e}")
