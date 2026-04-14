import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_guide():
    doc = docx.Document()
    
    # Main Document Title
    title = doc.add_heading('Guide d\'Utilisation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. App Short Description
    doc.add_heading('Description de l\'Application', level=1)
    doc.add_paragraph(
        "Procurement Watch (ou CapWatch) est une plateforme centralisée de veille et d'intelligence "
        "conçue pour traquer, analyser et gérer les opportunités et les appels d'offres internationaux. "
        "Elle extrait les données de multiples sources mondiales, intègre l'Intelligence Artificielle pour "
        "vérifier la pertinence des projets, et fournit à l'équipe un espace collaboratif pour la prise de décision."
    )
    
    # 2. Features breakdown
    doc.add_heading('Fonctionnalités', level=1)
    
    # Feature 1
    doc.add_heading('Tableau de Bord & Recherche Booléenne', level=2)
    doc.add_paragraph(
        "Utilisation : Depuis le tableau principal, vous pouvez visualiser tous les projets extraits. "
        "Utilisez la barre de recherche pour faire des requêtes simples ou complexes (ex: source:\"World Bank\" AND decision:Go). "
        "Vous pouvez également enregistrer vos recherches fréquentes et exporter les résultats sous format Excel."
    )

    # Feature 2
    doc.add_heading('Vue Chronologique (Timeline) & Analytique', level=2)
    doc.add_paragraph(
        "Utilisation : Cliquez sur l'onglet analytique pour observer les graphiques de performances et "
        "les tendances des projets. La Timeline regroupe les tâches et problèmes par semaine et par utilisateur, "
        "vous offrant une vue globale sur la charge de travail et la progression de l'équipe."
    )
    
    # Feature 3
    doc.add_heading('Actions Rapides depuis le Tableau', level=2)
    doc.add_paragraph(
        "Utilisation : Vous n'avez plus besoin d'ouvrir la page complète d'un projet pour chaque action. "
        "Modifiez directement le statut (transitions automatiques) ou éditez certains champs depuis les colonnes du tableau principal."
    )
    
    # Feature 4
    doc.add_heading('Inspecteur de Projet & Prise de Décision', level=2)
    doc.add_paragraph(
        "Utilisation : Cliquez sur une ligne de projet pour ouvrir le panneau latéral (ou cliquez pour voir la page complète dédiée). "
        "Vous pouvez y assigner des collègues, déclarer un projet \"Go\" ou \"No Go\", voter pour des opportunités prometteuses "
        "et modifier manuellement les dates limites."
    )
    
    # Feature 5
    doc.add_heading('Vérification IA et mode Deep Dive', level=2)
    doc.add_paragraph(
        "Utilisation : L'IA évalue automatiquement si un projet correspond à vos critères métiers (comme la cybersécurité). "
        "Depuis l'inspecteur de projet, cliquez sur \"Deep Dive\" pour que l'IA lance une recherche automatique des "
        "documents correspondants et résume les points clés de l'appel d'offres."
    )
    
    # Feature 6
    doc.add_heading('Discussions et Pièces jointes', level=2)
    doc.add_paragraph(
        "Utilisation : Au bas de chaque page de projet se trouve un fil de discussion. Tapez '@' pour mentionner "
        "et notifier un collègue. Vous pouvez joindre des images ou des documents (jusqu'à 20 Mo) qui peuvent souvent "
        "être ouverts ou prévisualisés directement dans l'application (ex: mode PDF viewer)."
    )
    
    # Feature 7
    doc.add_heading('Synchronisation des Données (Extraction)', level=2)
    doc.add_paragraph(
        "Utilisation : Via le bouton Sync en haut, vous pouvez forcer le robot à chercher de nouvelles données "
        "manuellement, et visualiser les résultats en temps réel. Vous pouvez également cliquer sur l'onglet Schedule "
        "pour planifier des routines (ex: tous les vendredis à 8h)."
    )
    
    # Feature 8
    doc.add_heading('Gestion d\'Équipe (Team Management)', level=2)
    doc.add_paragraph(
        "Utilisation : Accessibles via le menu d'administration, vous trouverez l'annuaire de votre équipe et "
        "les indicateurs de performance de chaque membre. Les administrateurs peuvent y créer des comptes, "
        "redéfinir les rôles et gérer (désactiver / supprimer) les accès très facilement."
    )
    
    # Feature 9
    doc.add_heading('Notes de Mise à Jour (Release Notes)', level=2)
    doc.add_paragraph(
        "Utilisation : Afin de suivre l'évolution de la plateforme, un pop-up et une page dédiée affichent les nouveautés. "
        "Elles sont classées en deux catégories claires : \"Nouvelles fonctionnalités\" et \"Corrections de bugs\"."
    )

    # Feature 10
    doc.add_heading('Gestion des Mots-Clés et Alertes', level=2)
    doc.add_paragraph(
        "Utilisation : Les paramètres de configuration vous permettent de définir quels termes l'application oit "
        "rechercher activement. Les utilisateurs reçoivent des alertes sonores ou des notifications lorsque de "
        "nouveaux projets correspondants sont trouvés, ou lorsqu'on les a mentionnés."
    )
    
    doc.save('Guide_Utilisateur_v2.docx')

if __name__ == '__main__':
    create_guide()
    print("Docx v2 created format applied.")
