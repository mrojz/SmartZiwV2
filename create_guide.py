import sys
import subprocess
import os

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install('python-docx')
    import docx

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading("Guide d'Utilisation Simplifié - Procurement Watch", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Define content structure
sections = [
    {
        "heading": "1. Introduction",
        "content": "Procurement Watch est une plateforme intelligente SaaS qui permet de suivre les opportunités d'appels d'offres internationaux à travers plusieurs sources. Elle automatise la recherche, vérification et l'analyse par Intelligence Artificielle (IA), et centralise les résultats dans un tableau de bord unique pour vos équipes."
    },
    {
        "heading": "2. Première Connexion",
        "content": "Lors de votre première connexion (avec les identifiants fournis par l'administrateur), il vous sera demandé de créer un nouveau mot de passe sécurisé pour protéger votre compte. Aucune inscription n'est publique, chaque utilisateur est créé par l'administrateur."
    },
    {
        "heading": "3. Tableau de Bord et Recherche",
        "content": "Le tableau de bord est votre espace de travail principal. Il affiche toutes les opportunités récupérées avec des informations telles que la source, la date limite, le verdict IA, les votes et les utilisateurs assignés.\n\n• Actions rapides : Vous pouvez trier les colonnes, utiliser des filtres (source, décision, conformité IA, continent) et faire des recherches textuelles rapides par mots-clés.\n• Recherche Avancée : Combinez plusieurs critères syntaxiques (ex: source:\"World Bank\" AND decision:Go) pour des requêtes précises.\n• Sauvegarde : Enregistrez jusqu'à 30 filtres de recherche préférés par utilisateur pour y accéder rapidement plus tard.\n• Exportation : Téléchargez l'intégralité ou une sélection filtrée de vos résultats au format Excel (.xlsx) d'un simple clic."
    },
    {
        "heading": "4. Inspecteur de Projet",
        "content": "En cliquant sur un projet dans le tableau de bord, un panneau latéral s'ouvre avec tous ses détails.\n\n• Informations clés : Consultez la date limite, le pays, le sponsor, les mots-clés correspondants, les documents liés trouvés par l'IA et l'analyse de pertinence détaillée.\n• Prise de décision : Les managers peuvent marquer un projet comme « Go » (À poursuivre) ou « No Go » (À ignorer).\n• Modification de la date limite : Si nécessaire pour une meilleure traçabilité, la date limite peut être modifiée manuellement par les managers/admins (l'ancienne est gardée en historique).\n• Votes : Tous les membres peuvent voter (👍 ou 👎) pour indiquer le potentiel d'un projet.\n• Assignation : Attribuez le projet à un ou plusieurs membres de votre équipe pour la revue croisée.\n• Deep Dive (Recherche approfondie IA) : Lancez une analyse IA poussée pour chercher des documents additionnels sur la page source et obtenir un résumé complet posté ensuite en tant que commentaire."
    },
    {
        "heading": "5. Discussions et Collaboration (Nouveau)",
        "content": "Chaque projet dispose de son propre fil de discussion pour fluidifier le travail d'équipe.\n\n• Commentaires : Échangez et ajoutez vos remarques (jusqu'à 4000 caractères par message).\n• Fichiers joints : Ajoutez des images (automatiquement compressées et affichées) ou des documents (visionneuse PDF intégrée dans l'app, téléchargement pour le reste, limite de 20 Mo par fichier).\n• Mentions : Utilisez @nom pour notifier un collègue spécifiquement, ce qui l'inscrit automatiquement aux notifications de ce projet."
    },
    {
        "heading": "6. Synchronisation des Données (Scraping)",
        "content": "C'est la fonctionnalité permettant à la plateforme de chercher de nouvelles offres sur 13+ sources compatibles (ex: Banque Mondiale, GIZ, IADB, DGMarket, IsDB, etc.).\n\n• Sync Manuelle : Depuis l'onglet « Sync », sélectionnez vos sources (vérification IA incluse ou non) et lancez à la demande.\n• Sync Planifiée : Automatisez les recherches (ex: fréquence journalière, hebdomadaire, etc.) via l'onglet « Schedule ».\n• Suivi & Historique : Vous pouvez consulter les journaux détaillés en temps réel, par scraper, et garder l'historique complet des exécutions, succès ou échecs."
    },
    {
        "heading": "7. Notifications",
        "content": "Restez à jour grâce aux notifications en temps réel de votre cloche en haut à droite :\n\n• Lorsqu'on vous mentionne (@) explicitement dans une discussion.\n• Lorsqu'un nouveau commentaire est publié sur un projet auquel vous êtes assigné ou inscrit.\n• Lorsqu'on vous assigne à une nouvelle opportunité.\n• Alertes sonores et notifications du navigateur à la découverte de nouveaux résultats suite à un scraping."
    },
    {
        "heading": "8. Configuration",
        "content": "Paramétrez le comportement global du système :\n\n• Mots-clés : Définissez les termes (français, anglais) devant obligatoirement apparaître lors du tri (ex: Cybersécurité, ISO 27001, VAPT).\n• Régions : Ajustez et gérez la répartition géographique (Pays, Régions géopolitiques, Continents) pour correspondre au mieux à vos critères de recherche organisationnels."
    },
    {
        "heading": "9. Gestion des Utilisateurs",
        "content": "Espace exclusif aux profils « Administrateurs ».\n\n• Création et Édition : Créez de nouveaux profils et assignez un niveau de droits adéquat (Admin, Manager, User).\n• Réinitialisation : Générez et imposez un nouveau mot de passe à un utilisateur.\n• Désactivation : Désactivez l'accès au compte avec préservation des données ou supprimez-le de façon permanente."
    },
    {
        "heading": "10. Notes de Mise à Jour (Release Notes) (Nouveau)",
        "content": "Les administrateurs peuvent ajouter depuis leurs panneaux des notes de version.\nEn retour, chaque utilisateur visualise, dès la connexion ou depuis le menu, les nouvelles mises à jour séparées de manière claire entre Nouvelles Fonctionnalités (New Features) et Corrections de Bugs (Bug Fixes)."
    }
]

for section in sections:
    doc.add_heading(section["heading"], level=1)
    
    paragraphs = section["content"].split('\n')
    for p in paragraphs:
        if p.strip() == '':
            continue
        if p.strip().startswith('•'):
            par = doc.add_paragraph(p.strip()[1:].strip(), style='List Bullet')
        else:
            par = doc.add_paragraph(p.strip())

output_path = r"d:\Dev\Ziw\new_cdx_gpt_5.4\Guide_Utilisateur_Procurement_Watch.docx"
doc.save(output_path)
print(f"Document saved to {output_path}")
